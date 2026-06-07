"""从PRD文档中提取所有模板的JSON字段配置，保存为template_fields.json"""
import json
import re
import sys
import os
from docx import Document

WENDANG_DIR = r"D:\LocalCodeProject\lvxiangbao\wendang"

# 所有包含JSON配置的文档
DOC_FILES = [
    "律享宝平台V1.2法律文书模板完整FreeMarker代码与前端表单JSON配置.docx",
    "律享宝平台V1.2法律文书模板完整FreeMarker代码与前端表单JSON配置（续）.docx",
]

def extract_json_configs(doc_path):
    """从一个docx中提取所有JSON字段配置"""
    configs = {}
    doc = Document(doc_path)
    
    for i, table in enumerate(doc.tables):
        raw = table.rows[0].cells[0].text.strip()
        
        # 跳过非JSON表格
        if not raw.startswith("JSON"):
            continue
        
        # 提取JSON部分
        text = re.sub(r'^JSON\s*', '', raw).strip()
        
        try:
            data = json.loads(text)
            tid = data.get("template_id", "")
            if tid and tid != "模板编号":  # 跳过表头示例
                configs[tid] = data
                print(f"  ✓ {tid}: {data.get('template_name', '')} - {len(data.get('fields', []))} fields")
        except json.JSONDecodeError as e:
            print(f"  ✗ Table {i}: JSON parse error - {e}")
    
    return configs

def extract_field_defs_from_prd(doc_path):
    """从PRD主文档提取字段定义表"""
    doc = Document(doc_path)
    field_defs = {}
    
    # Table 0: MS-TY通用文书(6份) - 5列: 编号,名称,必填字段,可选字段,特殊要求
    # Table 1: MS-AY案由文书(22份) - 3列: 编号,名称,核心必填字段
    # Tables 2-7: FZ法律援助文书 - 3列: 编号,名称,核心必填字段
    
    for tidx in range(min(8, len(doc.tables))):
        table = doc.tables[tidx]
        ncols = len(table.columns)
        
        for j, row in enumerate(table.rows):
            if j == 0:  # 跳过表头
                continue
            cells = [c.text.strip() for c in row.cells]
            tid = cells[0] if len(cells) > 0 else ""
            
            if not tid or not re.match(r'^(MS-|FZ-)', tid):
                continue
            
            field_defs[tid] = {
                "name": cells[1] if len(cells) > 1 else "",
                "required_fields": cells[2] if len(cells) > 2 else "",
                "optional_fields": cells[3] if len(cells) > 3 and ncols >= 4 else "",
                "special": cells[4] if len(cells) > 4 and ncols >= 5 else "",
            }
    
    return field_defs

def main():
    all_configs = {}
    
    # 从JSON配置文档提取
    for doc_file in DOC_FILES:
        doc_path = os.path.join(WENDANG_DIR, doc_file)
        if not os.path.exists(doc_path):
            print(f"跳过不存在的文件: {doc_file}")
            continue
        print(f"\n处理: {doc_file}")
        configs = extract_json_configs(doc_path)
        all_configs.update(configs)
    
    # 从PRD主文档提取字段定义
    prd_path = os.path.join(WENDANG_DIR, "律享宝律师工具箱平台V1.2迭代升级PRD需求文档.docx")
    if os.path.exists(prd_path):
        print(f"\n处理PRD主文档字段定义...")
        field_defs = extract_field_defs_from_prd(prd_path)
        print(f"  提取到 {len(field_defs)} 个模板的字段定义")
    else:
        field_defs = {}
    
    # 保存结果
    output = {
        "json_configs": all_configs,
        "prd_field_defs": field_defs,
        "summary": {
            "json_config_count": len(all_configs),
            "prd_field_def_count": len(field_defs),
            "covered_templates": sorted(set(list(all_configs.keys()) + list(field_defs.keys()))),
        }
    }
    
    out_path = os.path.join(os.path.dirname(__file__), "template_fields.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    
    print(f"\n=== 提取完成 ===")
    print(f"JSON配置: {len(all_configs)} 个模板")
    print(f"PRD字段定义: {len(field_defs)} 个模板")
    print(f"覆盖模板: {len(output['summary']['covered_templates'])} 个")
    print(f"输出文件: {out_path}")

if __name__ == "__main__":
    main()
