import docx
import json
import sys
import os

# Read the FreeMarker templates document
doc = docx.Document(r'D:\LocalCodeProject\lvxiangbao\wendang\律享宝平台V1.2法律文书模板完整FreeMarker代码与前端表单JSON配置.docx')

output_dir = r'D:\LocalCodeProject\lvxiangbao\backend\app\templates'
os.makedirs(output_dir, exist_ok=True)

# Extract all tables and save
all_data = []
for ti in range(len(doc.tables)):
    content = doc.tables[ti].rows[0].cells[0].text
    all_data.append({'table_idx': ti, 'content': content})

with open(os.path.join(output_dir, '_extracted_raw.txt'), 'w', encoding='utf-8') as f:
    for item in all_data:
        f.write(f"=== TABLE {item['table_idx']} ===\n")
        f.write(item['content'])
        f.write('\n\n')

# Also extract the "续" document
doc2 = docx.Document(r'D:\LocalCodeProject\lvxiangbao\wendang\律享宝平台V1.2法律文书模板完整FreeMarker代码与前端表单JSON配置（续）.docx')
all_data2 = []
for ti in range(len(doc2.tables)):
    content = doc2.tables[ti].rows[0].cells[0].text
    all_data2.append({'table_idx': ti, 'content': content})

with open(os.path.join(output_dir, '_extracted_raw_cont.txt'), 'w', encoding='utf-8') as f:
    for item in all_data2:
        f.write(f"=== TABLE {item['table_idx']} ===\n")
        f.write(item['content'])
        f.write('\n\n')

# Also extract the field definition document
doc3 = docx.Document(r'D:\LocalCodeProject\lvxiangbao\wendang\律享宝平台V1.2法律文书模板完整字段定义与FreeMarker占位符示例.docx')
with open(os.path.join(output_dir, '_field_defs.txt'), 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc3.paragraphs):
        if p.text.strip():
            f.write(f'P{i}: {p.text}\n')
    for ti, t in enumerate(doc3.tables):
        f.write(f'\n=== TABLE {ti} ({len(t.rows)}r x {len(t.columns)}c) ===\n')
        for ri, row in enumerate(t.rows):
            cells = [c.text[:200].replace('\n', ' ') for c in row.cells]
            f.write(f'  R{ri}: {cells}\n')

# Also extract the other field definition doc (1)
doc4 = docx.Document(r'D:\LocalCodeProject\lvxiangbao\wendang\律享宝平台V1.2法律文书模板完整字段定义与FreeMarker占位符示例 (1).docx')
with open(os.path.join(output_dir, '_field_defs_1.txt'), 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc4.paragraphs):
        if p.text.strip():
            f.write(f'P{i}: {p.text}\n')
    for ti, t in enumerate(doc4.tables):
        f.write(f'\n=== TABLE {ti} ({len(t.rows)}r x {len(t.columns)}c) ===\n')
        for ri, row in enumerate(t.rows):
            cells = [c.text[:200].replace('\n', ' ') for c in row.cells]
            f.write(f'  R{ri}: {cells}\n')

# Extract the W020240118789369618193.docx (official court template)
doc5 = docx.Document(r'D:\LocalCodeProject\lvxiangbao\wendang\W020240118789369618193.docx')
with open(os.path.join(output_dir, '_court_official.txt'), 'w', encoding='utf-8') as f:
    for i, p in enumerate(doc5.paragraphs):
        if p.text.strip():
            f.write(f'P{i}: {p.text}\n')
    for ti, t in enumerate(doc5.tables):
        f.write(f'\n=== TABLE {ti} ({len(t.rows)}r x {len(t.columns)}c) ===\n')
        for ri, row in enumerate(t.rows):
            cells = [c.text[:200].replace('\n', ' ') for c in row.cells]
            f.write(f'  R{ri}: {cells}\n')

print("All documents extracted successfully!")
print(f"Templates doc: {len(doc.tables)} tables")
print(f"Templates cont doc: {len(doc2.tables)} tables")
print(f"Field defs doc: {len(doc3.paragraphs)} paras, {len(doc3.tables)} tables")
print(f"Field defs (1) doc: {len(doc4.paragraphs)} paras, {len(doc4.tables)} tables")
print(f"Court official doc: {len(doc5.paragraphs)} paras, {len(doc5.tables)} tables")
